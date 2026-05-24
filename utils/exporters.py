"""
utils/exporters.py
Convert markdown investment notes to .txt, .docx, and .pdf for download.

Dependencies:
    pip install python-docx reportlab
"""

import io
import re


def export_to_txt(md_text: str) -> bytes:
    """Strips markdown to clean plain text."""
    lines = []
    for line in md_text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('### '):
            lines.append('')
            lines.append(stripped[4:].upper())
            lines.append('-' * len(stripped[4:]))
        elif stripped.startswith('## '):
            lines.append('')
            lines.append(stripped[3:].upper())
            lines.append('=' * len(stripped[3:]))
        elif stripped.startswith('# '):
            lines.append(stripped[2:].upper())
        else:
            cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)
            lines.append(cleaned)
    return '\n'.join(lines).encode('utf-8')


def export_to_docx(md_text: str, company_name: str = "Company") -> bytes:
    """Converts markdown to a styled Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    header = doc.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run('Pivot & Co · Agricultural Impact Fund')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith(('- ', '* ')):
            _add_formatted_paragraph(doc, stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            _add_formatted_paragraph(doc, text, style='List Number')
        else:
            _add_formatted_paragraph(doc, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_paragraph(doc, text: str, style: str = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def export_to_pdf(md_text: str, company_name: str = "Company") -> bytes:
    """Converts markdown to a styled PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('H1', parent=styles['Heading1'], textColor=HexColor('#166534'),
                        spaceAfter=10, fontSize=18)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], textColor=HexColor('#166534'),
                        spaceAfter=6, spaceBefore=12, fontSize=13)
    body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=6)
    bullet = ParagraphStyle('Bullet', parent=body, leftIndent=18)
    footer = ParagraphStyle('Footer', parent=styles['Italic'], textColor=HexColor('#6B7280'),
                            fontSize=8, alignment=2)

    story = []
    story.append(Paragraph("Pivot &amp; Co · Agricultural Impact Fund", footer))
    story.append(Spacer(1, 0.15*inch))

    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.08*inch))
            continue
        if stripped.startswith('### '):
            story.append(Paragraph(_md_to_html(stripped[4:]), h2))
        elif stripped.startswith('## '):
            story.append(Paragraph(_md_to_html(stripped[3:]), h1))
        elif stripped.startswith('# '):
            story.append(Paragraph(_md_to_html(stripped[2:]), h1))
        elif stripped.startswith(('- ', '* ')):
            story.append(Paragraph(f'• {_md_to_html(stripped[2:])}', bullet))
        elif re.match(r'^\d+\.\s', stripped):
            story.append(Paragraph(_md_to_html(stripped), bullet))
        else:
            story.append(Paragraph(_md_to_html(stripped), body))

    doc.build(story)
    return buffer.getvalue()


def _md_to_html(text: str) -> str:
    text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', text)
    return text